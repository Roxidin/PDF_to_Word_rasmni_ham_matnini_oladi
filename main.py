import os
from pdf2docx import Converter
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from getpass import getpass
import pikepdf
import tempfile

# Tesseract yo'lini sozlash
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def remove_pdf_password(pdf_file, password):
    """PDF faylning parolini olib tashlaydi."""
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        temp_file_path = temp_file.name
        temp_file.close()
        with pikepdf.open(pdf_file, password=password) as pdf:
            pdf.save(temp_file_path)
        print(f"Parol olib tashlandi: {temp_file_path}")
        return temp_file_path
    except pikepdf.PasswordError:
        raise ValueError("Noto'g'ri parol.")
    except Exception as e:
        raise Exception(f"Xatolik: {str(e)}")

def pdf_to_word(pdf_file, word_file, password=None, max_attempts=3):
    """PDF faylni Word fayliga aylantiradi (rasmlardagi matnni ham o'qiydi)."""
    temp_pdf_file = None
    attempts = 0
    while attempts < max_attempts:
        try:
            if not os.path.exists(pdf_file):
                raise FileNotFoundError(f"{pdf_file} topilmadi.")

            # 1-qadam: pdf2docx bilan sinab ko'ramiz
            converter = Converter(pdf_file, password=password)
            converter.convert(word_file, start=0, end=None)
            converter.close()

            # Matn bor-yo'qligini tekshirish
            doc = Document(word_file)
            has_text = any(para.text.strip() for para in doc.paragraphs)

            if has_text:
                print(f"{pdf_file} Word fayliga aylantirildi (matn sifatida).")
                return True
            else:
                print("Matn topilmadi, OCR boshlanmoqda...")

                # 2-qadam: Parolni olib tashlash
                if password:
                    temp_pdf_file = remove_pdf_password(pdf_file, password)
                    pdf_file_to_process = temp_pdf_file
                else:
                    pdf_file_to_process = pdf_file

                # 3-qadam: OCR jarayoni
                poppler_path = r'C:\Program Files\poppler-24.08.0\Library\bin'
                images = convert_from_path(pdf_file_to_process, poppler_path=poppler_path)

                doc = Document()
                for i, image in enumerate(images):
                    text = pytesseract.image_to_string(image, lang='uzb+rus')
                    doc.add_paragraph(text)
                    doc.add_paragraph(f"--- Sahifa {i+1} ---")

                doc.save(word_file)
                print(f"{pdf_file} OCR yordamida Word fayliga aylantirildi.")
                return True

        except Exception as e:
            if "password" in str(e).lower() and password is None:
                print("PDF parol bilan himoyalangan.")
                password = input(f"Parolni kiriting (urinish {attempts + 1}/{max_attempts}): ")
                attempts += 1
            elif "password" in str(e).lower() and password is not None:
                print(f"Noto'g'ri parol. Qayta urinib ko'ring (urinish {attempts + 1}/{max_attempts}).")
                password = input("Parolni qayta kiriting: ")
                attempts += 1
            else:
                print(f"Xatolik: {str(e)}")
                return False
        finally:
            if temp_pdf_file and os.path.exists(temp_pdf_file):
                os.remove(temp_pdf_file)

    print(f"Parol urinishlari tugadi ({max_attempts}).")
    return False

if __name__ == '__main__':
    pdf_path = r"C:\Users\GENERAL\Desktop\Chaqiruv_qog`ozi-343231104644.pdf"
    word_path = r"C:\Users\GENERAL\Desktop\Chaqiruv_qog`ozi-343231104644.docx"

    pdf_password = "2025"  # To'g'ri parolni yozing
    pdf_to_word(pdf_path, word_path, password=pdf_password)